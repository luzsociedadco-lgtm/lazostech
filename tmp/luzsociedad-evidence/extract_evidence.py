from __future__ import annotations

import json
import os
import shutil
import sys
import zipfile
from pathlib import Path

import pdfplumber
from docx import Document
from docx.oxml.ns import qn


def extract_docx(path: Path, out_dir: Path) -> dict:
    doc = Document(path)
    media_dir = out_dir / f"{path.stem}_media"
    media_dir.mkdir(parents=True, exist_ok=True)

    with zipfile.ZipFile(path) as archive:
        for name in archive.namelist():
            if name.startswith("word/media/") and not name.endswith("/"):
                target = media_dir / Path(name).name
                with archive.open(name) as source, target.open("wb") as sink:
                    shutil.copyfileobj(source, sink)

    paragraphs = []
    for index, paragraph in enumerate(doc.paragraphs):
        image_targets = []
        for blip in paragraph._p.xpath(".//a:blip"):
            rel_id = blip.get(qn("r:embed"))
            if rel_id and rel_id in doc.part.rels:
                image_targets.append(Path(doc.part.rels[rel_id].target_ref).name)
        text = paragraph.text.strip()
        if text or image_targets:
            paragraphs.append(
                {
                    "index": index,
                    "text": text,
                    "style": paragraph.style.name if paragraph.style else "",
                    "images": image_targets,
                }
            )

    tables = []
    for table_index, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append({"index": table_index, "rows": rows})

    return {
        "type": "docx",
        "path": str(path),
        "paragraphs": paragraphs,
        "tables": tables,
        "media_dir": str(media_dir),
        "media_files": sorted(file.name for file in media_dir.iterdir()),
    }


def extract_pdf(path: Path) -> dict:
    pages = []
    with pdfplumber.open(path) as pdf:
        for index, page in enumerate(pdf.pages, start=1):
            pages.append(
                {
                    "page": index,
                    "text": (page.extract_text() or "").strip(),
                    "width": page.width,
                    "height": page.height,
                    "image_count": len(page.images),
                }
            )
    return {"type": "pdf", "path": str(path), "pages": pages}


def main() -> None:
    out_dir = Path(sys.argv[1])
    out_dir.mkdir(parents=True, exist_ok=True)
    paths = [Path(value) for value in sys.argv[2:]]
    results = []
    for path in paths:
        if path.suffix.lower() == ".docx":
            results.append(extract_docx(path, out_dir))
        elif path.suffix.lower() == ".pdf":
            results.append(extract_pdf(path))
    output_path = out_dir / "evidence.json"
    output_path.write_text(json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8")
    print(output_path)


if __name__ == "__main__":
    main()
